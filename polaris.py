from flask import Flask, render_template_string, jsonify, request, send_file
from flask_socketio import SocketIO, emit
import requests
import base64
import time
import os
import json
import threading
import hashlib
import concurrent.futures
import asyncio
from dotenv import load_dotenv
from pyngrok import ngrok
import PyPDF2
import docx
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from werkzeug.utils import secure_filename
import io
import re
from collections import defaultdict
from datetime import datetime, date
from fpdf import FPDF
import random

# Load environment variables
load_dotenv()

# API Keys
CONFIG_FILE = "user_config.json"

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {}

def save_config(config):
    with open(CONFIG_FILE, "w") as f:
        json.dump(config, f, indent=2)

config = load_config()

RECALL_API_KEY = config.get("RECALL_API_KEY")
CEREBRAS_API_KEY = config.get("CEREBRAS_API_KEY")
MURF_API_KEY = config.get("MURF_API_KEY")

app = Flask(__name__)
app.config['SECRET_KEY'] = 'polaris-meeting-bot'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

@app.route("/get_api_keys")
def get_api_keys():
    return jsonify({
        "RECALL_API_KEY": bool(RECALL_API_KEY),
        "CEREBRAS_API_KEY": bool(CEREBRAS_API_KEY),
        "MURF_API_KEY": bool(MURF_API_KEY),
    })

@app.route("/set_api_keys", methods=["POST"])
def set_api_keys():
    global RECALL_API_KEY, CEREBRAS_API_KEY, MURF_API_KEY
    data = request.json

    if "RECALL_API_KEY" in data:
        RECALL_API_KEY = data["RECALL_API_KEY"]
    if "CEREBRAS_API_KEY" in data:
        CEREBRAS_API_KEY = data["CEREBRAS_API_KEY"]
    if "MURF_API_KEY" in data:
        MURF_API_KEY = data["MURF_API_KEY"]

    save_config({
        "RECALL_API_KEY": RECALL_API_KEY,
        "CEREBRAS_API_KEY": CEREBRAS_API_KEY,
        "MURF_API_KEY": MURF_API_KEY,
    })

    return jsonify({"success": True})

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
socketio = SocketIO(app, cors_allowed_origins="*")

# Global state
tunnel_url = None
current_bot_id = None
transcript_buffer = ""
last_update_time = 0
sentence_timeout_handle = None
audio_cache = {}

# NEW: Mute/Unmute functionality
polaris_muted = False

# In-call memory and conversation tracking
conversation_memory = []
full_meeting_transcript = []
meeting_start_time = None
generated_mom_file = None

# Bot startup state management
bot_startup_state = {
    'is_joining': False,
    'join_start_time': None,
    'is_ready': False,
    'countdown_timer': None,
    'greeting_sent': False
}

# Response settings
response_settings = {
    'mode': 'detailed',
    'max_tokens': 70,
    'response_style': 'helpful and contextual'
}

voice_options = [
    ("Marcus (US Male)", "en-US-marcus"),
    ("Natalie (US Female)", "en-US-natalie"),
    ("Amara (US Female)", "en-US-amara"),
    ("Charles (US Male)", "en-US-charles"),
    ("Freddie (UK Male)", "en-UK-freddie"),
    ("Emma (UK Female)", "en-UK-emma"),
    ("Oliver (UK Male)", "en-UK-oliver"),
    ("Sarah (US Female)", "en-US-sarah"),
    ("David (UK Male)", "en-UK-david"),
    ("Sophie (UK Female)", "en-UK-sophie"),
    ("Alex (US Male)", "en-US-alex"),
    ("Lily (US Female)", "en-US-lily"),
    ("James (UK Male)", "en-UK-james"),
    ("Grace (UK Female)", "en-UK-grace"),
    ("Ryan (US Male)", "en-US-ryan")
]

selected_voice_id = "en-US-marcus"  # Default value

# MOM generation functions
def prepare_mom_prompt_exact(full_transcript):
    today = date.today().strftime("%Y-%m-%d")
    transcript_text = "\n".join([entry['text'] for entry in full_transcript])

    date_pattern = re.compile(r'\b(\d{4}-\d{2}-\d{2})\b')
    match = date_pattern.search(transcript_text)
    if match:
        date_found = match.group(1)
    else:
        date_found = today

    prompt = (
        f"Prepare a MOM (Minutes of Meeting) for the following transcript in a suitable format. "
        f"Date: {date_found}. "
        f"Do not produce any pre- or post-texts. Generate insights based on the content logically. Keep it as concise as possible. "
        f"Remove all unnecessary symbols or formatting: {transcript_text}"
    )
    return prompt

def generate_mom_pdf_exact(mom_content, filename="meeting_minutes.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    cleaned_text = mom_content.encode('latin-1', 'ignore').decode('latin-1')

    if cleaned_text.strip():
        pdf.multi_cell(0, 10, cleaned_text)
    else:
        pdf.multi_cell(0, 10, "Error: No valid content to display.")

    pdf_bytes = pdf.output(dest="S").encode('latin-1')
    return pdf_bytes

async def generate_meeting_summary_fixed():
    global full_meeting_transcript, generated_mom_file

    if not full_meeting_transcript:
        print("‚ùå No meeting transcript to summarize")
        return None, None

    try:
        print("üìù Generating MOM...")
        mom_prompt = prepare_mom_prompt_exact(full_meeting_transcript)

        def generate_mom_exact():
            headers = {
                'Authorization': f'Bearer {CEREBRAS_API_KEY}',
                'Content-Type': 'application/json'
            }

            payload = {
                'model': 'llama3.1-8b',
                'messages': [
                    {
                        'role': 'system',
                        'content': 'You are a meeting secretary. Follow the exact format requested. Do not add extra information not in the transcript.'
                    },
                    {
                        'role': 'user',
                        'content': mom_prompt
                    }
                ],
                'max_tokens': 1000,
                'temperature': 0.2,
                'top_p': 0.9
            }

            response = session.post(
                'https://api.cerebras.ai/v1/chat/completions',
                json=payload,
                headers=headers,
                timeout=12
            )

            if response.status_code == 200:
                result = response.json()
                return result['choices'][0]['message']['content'].strip()
            else:
                print(f"‚ùå MOM generation error: {response.status_code}")
                return None

        loop = asyncio.get_event_loop()
        executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
        mom_content = await loop.run_in_executor(executor, generate_mom_exact)

        if mom_content:
            pdf_bytes = generate_mom_pdf_exact(mom_content)
            if pdf_bytes:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                pdf_filename = f"meeting_minutes_{timestamp}.pdf"
                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)

                with open(pdf_path, 'wb') as f:
                    f.write(pdf_bytes)

                generated_mom_file = pdf_filename
                print(f"‚úÖ MOM generated: {pdf_filename}")
                return mom_content, pdf_filename

        return None, None

    except Exception as e:
        print(f"‚ùå MOM generation error: {e}")
        return None, None

# Bot leave function
def leave_bot():
    global current_bot_id, bot_startup_state, conversation_memory, full_meeting_transcript

    if not current_bot_id:
        print("‚ùå No active bot to leave")
        return False, "No active bot found"

    try:
        url = f"https://us-west-2.recall.ai/api/v1/bot/{current_bot_id}/leave_call/"
        headers = {
            'Authorization': f'Token {RECALL_API_KEY}',
            'Accept': 'application/json'
        }

        response = session.post(url, headers=headers, timeout=8)

        if response.status_code == 200:
            print(f"‚úÖ Bot {current_bot_id} left meeting successfully")

            def generate_summary_async():
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                mom_content, pdf_filename = loop.run_until_complete(generate_meeting_summary_fixed())
                if mom_content and pdf_filename:
                    socketio.emit('mom_generated', {
                        'content': mom_content[:500] + "..." if len(mom_content) > 500 else mom_content,
                        'pdf_filename': pdf_filename,
                        'message': "üìù MOM generated! Download available."
                    })
                loop.close()

            threading.Thread(target=generate_summary_async, daemon=True).start()

            current_bot_id = None
            bot_startup_state = {
                'is_joining': False,
                'join_start_time': None,
                'is_ready': False,
                'countdown_timer': None,
                'greeting_sent': False
            }

            conversation_memory = []
            return True, "Bot left meeting successfully"
        else:
            print(f"‚ùå Failed to leave meeting: {response.text}")
            return False, f"Leave failed: {response.text}"

    except Exception as e:
        print(f"‚ùå Error leaving bot: {e}")
        return False, str(e)

def check_polaris_mention(text):
    """IMPROVED: Better Polaris detection with context awareness"""
    if not text or not text.strip():
        return False

    words = text.strip().split()
    first_three = words[:3] if len(words) >= 3 else words

    for word in first_three:
        clean_word = re.sub(r'[^\w]', '', word.lower())
        if clean_word == 'polaris':
            return True

    return False

# FULLY FIXED RAG SYSTEM - COMBINING BEST OF BOTH
class DocumentProcessor:
    def __init__(self):
        print("üß≠ Initializing COMPLETELY FIXED RAG system...")

        # Initialize sentence transformer with optimal settings
        self.encoder = SentenceTransformer('all-MiniLM-L6-v2', device='cpu')
        self.encoder.max_seq_length = 384  # Increased for better context

        # CRITICAL FIX: Use L2 distance for proper cosine similarity
        self.dimension = 384
        self.index = faiss.IndexFlatL2(self.dimension)  # FIXED: L2 instead of IP

        # Document storage
        self.documents = []
        self.document_metadata = []
        self.document_files = {}
        self.search_cache = {}

        # Enhanced text processing
        self.sentence_splitter = re.compile(r'[.!?]+\s+|\n\s*\n')
        self.cleanup_regex = re.compile(r'\s+')

        print("‚úÖ COMPLETELY FIXED RAG system ready!")

    def clean_text(self, text):
        """Enhanced text cleaning"""
        if not text:
            return ""

        # Remove non-printable characters
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', ' ', text)

        # Normalize whitespace
        text = self.cleanup_regex.sub(' ', text)

        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)

        return text.strip()

    def split_text(self, text, max_length=500, overlap=80):
        """FIXED: Enhanced text splitting with better overlap"""
        if not text or not text.strip():
            return []

        text = self.clean_text(text)
        if not text:
            return []

        chunks = []

        # Split by sentences and paragraphs
        sentences = self.sentence_splitter.split(text)
        sentences = [s.strip() for s in sentences if s.strip()]

        if not sentences:
            return []

        current_chunk = ""

        for sentence in sentences:
            # If adding this sentence would exceed max_length
            if len(current_chunk) + len(sentence) + 1 > max_length:
                if current_chunk:
                    chunks.append(current_chunk.strip())

                    # Create overlap from the end of current chunk
                    words = current_chunk.split()
                    overlap_words = words[-min(overlap//6, len(words)):]
                    current_chunk = " ".join(overlap_words)

                    # If the sentence itself is very long, split it
                    if len(sentence) > max_length:
                        sentence_words = sentence.split()
                        for i in range(0, len(sentence_words), max_length//10):
                            chunk_words = sentence_words[i:i + max_length//10]
                            chunk_text = " ".join(chunk_words)
                            if chunk_text.strip():
                                chunks.append(chunk_text.strip())
                        current_chunk = ""
                    else:
                        current_chunk = (current_chunk + " " + sentence).strip()
                else:
                    # Very long sentence, split it
                    if len(sentence) > max_length:
                        words = sentence.split()
                        for i in range(0, len(words), max_length//10):
                            chunk_words = words[i:i + max_length//10]
                            chunk_text = " ".join(chunk_words)
                            if chunk_text.strip():
                                chunks.append(chunk_text.strip())
                    else:
                        current_chunk = sentence
            else:
                current_chunk = (current_chunk + " " + sentence).strip() if current_chunk else sentence

        # Add the final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # Enhanced filtering - keep chunks with meaningful content
        filtered_chunks = []
        for chunk in chunks:
            # Remove chunks that are too short or just punctuation/numbers
            if len(chunk.strip()) >= 50 and re.search(r'[a-zA-Z]{3,}', chunk):
                filtered_chunks.append(chunk)

        print(f"üìÑ Text chunking: {len(sentences)} sentences ‚Üí {len(chunks)} chunks ‚Üí {len(filtered_chunks)} filtered chunks")
        return filtered_chunks

    def process_pdf(self, file_path):
        """ENHANCED PDF processing with better error handling"""
        try:
            chunks = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            cleaned_text = self.clean_text(page_text)
                            if cleaned_text:
                                text += cleaned_text + "\n\n"
                    except Exception as e:
                        print(f"‚ö†Ô∏è Warning: Error processing page {page_num + 1}: {e}")
                        continue

            if not text.strip():
                print("‚ùå No text extracted from PDF")
                return []

            chunks = self.split_text(text, max_length=500, overlap=80)
            print(f"üìÑ PDF processed: {len(pdf_reader.pages)} pages ‚Üí {len(text)} chars ‚Üí {len(chunks)} chunks")
            return chunks

        except Exception as e:
            print(f"‚ùå PDF processing error: {e}")
            return []

    def process_docx(self, file_path):
        """ENHANCED DOCX processing"""
        try:
            doc = docx.Document(file_path)
            text = ""

            # Process paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    cleaned_para = self.clean_text(para_text)
                    if cleaned_para:
                        text += cleaned_para + "\n\n"

            # Process tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        cleaned_row = self.clean_text(row_text)
                        if cleaned_row:
                            text += cleaned_row + "\n"
                text += "\n"

            if not text.strip():
                print("‚ùå No text extracted from DOCX")
                return []

            chunks = self.split_text(text, max_length=500, overlap=80)
            print(f"üìÑ DOCX processed: {len(text)} chars ‚Üí {len(chunks)} chunks")
            return chunks

        except Exception as e:
            print(f"‚ùå DOCX processing error: {e}")
            return []

    def process_txt(self, file_path):
        """ENHANCED TXT processing with encoding detection"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    print(f"‚úÖ Successfully read file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                print("‚ùå Could not read file with any encoding")
                return []

            text = self.clean_text(text)
            if not text:
                print("‚ùå No valid text after cleaning")
                return []

            chunks = self.split_text(text, max_length=500, overlap=80)
            print(f"üìÑ TXT processed: {len(text)} chars ‚Üí {len(chunks)} chunks")
            return chunks

        except Exception as e:
            print(f"‚ùå TXT processing error: {e}")
            return []

    def add_document(self, filename, chunks):
        """FIXED: Enhanced document addition with proper validation"""
        try:
            print(f"üìö Processing {len(chunks)} chunks from {filename}...")

            if not chunks:
                print("‚ùå No chunks to add")
                return False

            # Validate chunks
            valid_chunks = []
            for chunk in chunks:
                if isinstance(chunk, str) and len(chunk.strip()) >= 50:
                    valid_chunks.append(chunk.strip())

            if not valid_chunks:
                print("‚ùå No valid chunks after filtering")
                return False

            print(f"üìä Valid chunks: {len(valid_chunks)} out of {len(chunks)}")

            # CRITICAL FIX: Generate embeddings with proper normalization
            try:
                embeddings = self.encoder.encode(
                    valid_chunks, 
                    batch_size=32,
                    show_progress_bar=False,
                    convert_to_tensor=False
                )

                # FIXED: Proper normalization for L2 distance
                embeddings = embeddings.astype(np.float32)
                norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
                norms[norms == 0] = 1  # Avoid division by zero
                embeddings = embeddings / norms

                print(f"üìä Generated embeddings shape: {embeddings.shape}")

            except Exception as e:
                print(f"‚ùå Embedding generation error: {e}")
                return False

            # Add to FAISS index
            try:
                self.index.add(embeddings)
                print(f"üìä Added embeddings to FAISS index. Total vectors: {self.index.ntotal}")
            except Exception as e:
                print(f"‚ùå FAISS index error: {e}")
                return False

            # Update document storage
            start_idx = len(self.documents)
            for i, chunk in enumerate(valid_chunks):
                self.documents.append(chunk)
                self.document_metadata.append({
                    'filename': filename,
                    'chunk_id': len(self.documents) - 1,
                    'chunk_index': i,
                    'length': len(chunk),
                    'content_preview': chunk[:100] + "..." if len(chunk) > 100 else chunk,
                    'word_count': len(chunk.split())
                })

            self.document_files[filename] = {
                'chunk_count': len(valid_chunks),
                'chunk_indices': list(range(start_idx, start_idx + len(valid_chunks))),
                'upload_time': datetime.now().isoformat(),
                'total_chars': sum(len(chunk) for chunk in valid_chunks),
                'avg_chunk_length': sum(len(chunk) for chunk in valid_chunks) // len(valid_chunks)
            }

            # Clear search cache
            self.search_cache.clear()
            print(f"‚úÖ Added {len(valid_chunks)} chunks from {filename}")
            print(f"üìä Total documents: {len(self.document_files)}, Total chunks: {len(self.documents)}")
            return True

        except Exception as e:
            print(f"‚ùå Document addition error: {e}")
            return False

    def search(self, query, top_k=5):
        """COMPLETELY FIXED: Proper search with L2 distance and cosine similarity"""
        try:
            # Input validation
            if not query or not query.strip():
                print("‚ùå Empty query")
                return []

            query = query.strip()
            cache_key = hashlib.md5(f"{query}_{top_k}".encode()).hexdigest()

            if cache_key in self.search_cache:
                print(f"‚ö° Cache hit for query: '{query[:50]}...'")
                return self.search_cache[cache_key]

            if len(self.documents) == 0:
                print("‚ùå No documents in RAG system")
                return []

            print(f"üîç Searching {len(self.documents)} chunks for: '{query[:50]}...'")

            # Generate query embedding
            try:
                query_embedding = self.encoder.encode(
                    [query], 
                    show_progress_bar=False,
                    convert_to_tensor=False
                )
                query_embedding = query_embedding.astype(np.float32)

                # Normalize for cosine similarity
                norm = np.linalg.norm(query_embedding, axis=1, keepdims=True)
                if norm[0] > 0:
                    query_embedding = query_embedding / norm
                else:
                    print("‚ùå Query embedding has zero norm")
                    return []

            except Exception as e:
                print(f"‚ùå Query embedding error: {e}")
                return []

            # FIXED: Search with better parameters
            search_k = min(top_k * 2, len(self.documents))

            try:
                # L2 distance search
                distances, indices = self.index.search(query_embedding, search_k)

                # Convert L2 distances to cosine similarities
                # For normalized vectors: cosine_sim = 1 - (L2_dist^2 / 2)
                similarities = 1 - (distances[0] ** 2) / 2

            except Exception as e:
                print(f"‚ùå FAISS search error: {e}")
                return []

            results = []
            for similarity, idx in zip(similarities, indices[0]):
                if idx < len(self.documents):
                    # CRITICAL FIX: Much lower threshold for better recall
                    if similarity > 0.1:  # Lowered from 0.3 to 0.1
                        chunk_text = self.documents[idx]

                        # Additional relevance scoring
                        query_words = set(query.lower().split())
                        chunk_words = set(chunk_text.lower().split())
                        word_overlap = len(query_words.intersection(chunk_words))

                        # Boost score if there's word overlap
                        adjusted_score = similarity
                        if word_overlap > 0:
                            boost = (word_overlap / len(query_words)) * 0.2
                            adjusted_score += boost

                        results.append({
                            'chunk': chunk_text,
                            'metadata': self.document_metadata[idx],
                            'score': float(adjusted_score),
                            'similarity': float(similarity),
                            'word_overlap': word_overlap
                        })

            # Sort by adjusted score
            results.sort(key=lambda x: x['score'], reverse=True)
            results = results[:top_k]

            self.search_cache[cache_key] = results

            if results:
                scores = [r['score'] for r in results]
                print(f"‚úÖ Found {len(results)} relevant chunks with scores: {scores}")
                for i, result in enumerate(results[:3]):
                    print(f"   {i+1}. Score: {result['score']:.3f} | {result['chunk'][:100]}...")
            else:
                print(f"‚ùå No relevant chunks found for query: '{query}'")
                print(f"üìä Total indexed chunks: {len(self.documents)}")

                # Debug: Show top similarities even if below threshold
                if len(similarities) > 0:
                    max_sim = max(similarities)
                    print(f"üìä Highest similarity found: {max_sim:.3f}")

            return results

        except Exception as e:
            print(f"‚ùå Search error: {e}")
            import traceback
            traceback.print_exc()
            return []

    def remove_document(self, filename):
        """Enhanced document removal"""
        try:
            if filename not in self.document_files:
                return False, "Document not found"

            print(f"üóëÔ∏è Removing document: {filename}")

            chunk_indices = set(self.document_files[filename]['chunk_indices'])

            new_documents = []
            new_metadata = []

            for i, (doc, meta) in enumerate(zip(self.documents, self.document_metadata)):
                if i not in chunk_indices:
                    new_documents.append(doc)
                    new_metadata.append({
                        **meta,
                        'chunk_id': len(new_documents) - 1
                    })

            # Rebuild the index
            if new_documents:
                embeddings = self.encoder.encode(
                    new_documents, 
                    batch_size=32, 
                    show_progress_bar=False,
                    convert_to_tensor=False
                )
                embeddings = embeddings.astype(np.float32)
                norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
                norms[norms == 0] = 1
                embeddings = embeddings / norms

                self.index = faiss.IndexFlatL2(self.dimension)
                self.index.add(embeddings)
            else:
                self.index = faiss.IndexFlatL2(self.dimension)

            self.documents = new_documents
            self.document_metadata = new_metadata

            # Update remaining file indices
            remaining_files = {}
            current_idx = 0
            for fname, finfo in self.document_files.items():
                if fname != filename:
                    chunk_count = finfo['chunk_count']
                    remaining_files[fname] = {
                        **finfo,
                        'chunk_indices': list(range(current_idx, current_idx + chunk_count))
                    }
                    current_idx += chunk_count

            self.document_files = remaining_files
            self.search_cache.clear()

            print(f"‚úÖ Removed document: {filename}")
            return True, "Document removed successfully"

        except Exception as e:
            print(f"‚ùå Document removal error: {e}")
            return False, str(e)

    def get_document_list(self):
        """Get list of uploaded documents with stats"""
        documents = []
        for filename, info in self.document_files.items():
            documents.append({
                'filename': filename,
                'chunk_count': info['chunk_count'],
                'upload_time': info['upload_time'],
                'total_chars': info.get('total_chars', 0),
                'avg_chunk_length': info.get('avg_chunk_length', 0)
            })
        return documents

    def get_stats(self):
        """Get detailed RAG system statistics"""
        return {
            'total_chunks': len(self.documents),
            'total_documents': len(self.document_files),
            'vector_dimension': self.dimension,
            'cache_size': len(self.search_cache),
            'index_size': self.index.ntotal if hasattr(self.index, 'ntotal') else 0,
            'avg_chunk_length': sum(len(doc) for doc in self.documents) // len(self.documents) if self.documents else 0
        }

# Initialize FIXED RAG system
rag_system = DocumentProcessor()

# Session setup with connection pooling
session = requests.Session()
session.headers.update({
    'User-Agent': 'Polaris-AI/1.0',
    'Connection': 'keep-alive',
    'Keep-Alive': 'timeout=20, max=50'
})

adapter = requests.adapters.HTTPAdapter(
    pool_connections=10,
    pool_maxsize=20,
    max_retries=0,
    pool_block=False
)
session.mount('http://', adapter)
session.mount('https://', adapter)

def create_recall_bot_correct(meeting_url, webhook_url):
    url = "https://us-west-2.recall.ai/api/v1/bot/"

    headers = {
        'Authorization': f'Token {RECALL_API_KEY}',
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }

    payload = {
        "meeting_url": meeting_url,
        "bot_name": "Polaris AI Assistant",
        "recording_config": {
            "transcript": {
                "provider": {
                    "meeting_captions": {}
                }
            },
            "realtime_endpoints": [
                {
                    "type": "webhook",
                    "url": webhook_url,
                    "events": ["transcript.data"]
                }
            ],
            "video_mixed_layout": "audio_only",
            "include_bot_in_recording": {
                "audio": False,
                "video": False
            },
            "start_recording_on": "participant_join"
        }
    }

    try:
        print(f"üß≠ Creating Polaris bot...")
        response = session.post(url, headers=headers, json=payload, timeout=10)

        if response.status_code == 201:
            result = response.json()
            bot_id = result.get('id')
            print(f"‚úÖ Polaris bot created: {bot_id}")
            return bot_id, None
        else:
            error_text = response.text
            print(f"‚ùå Error {response.status_code}: {error_text}")
            return None, f"Status {response.status_code}: {error_text}"

    except Exception as e:
        print(f"‚ùå Exception: {e}")
        return None, str(e)

def start_bot_countdown():
    """Random countdown 16-19 seconds, start button after 5 seconds"""
    global bot_startup_state, meeting_start_time, conversation_memory, full_meeting_transcript

    meeting_start_time = datetime.now()
    conversation_memory = []
    full_meeting_transcript = []

    bot_startup_state['is_joining'] = True
    bot_startup_state['join_start_time'] = time.time()
    bot_startup_state['is_ready'] = False
    bot_startup_state['greeting_sent'] = False

    countdown_duration = random.randint(16, 19)
    print(f"‚è∞ Starting {countdown_duration}-second countdown...")

    def countdown_thread():
        for i in range(countdown_duration, 0, -1):
            if not bot_startup_state['is_joining']:
                return

            socketio.emit('countdown_update', {
                'seconds_remaining': i,
                'message': f"üß≠ Polaris joining in {i} seconds..."
            })
            print(f"‚è∞ Polaris joining in {i} seconds...")

            # Enable start button after 5 seconds
            if i == countdown_duration - 20:
                socketio.emit('enable_start_button', {
                })

            time.sleep(1)

        if bot_startup_state['is_joining']:
            socketio.emit('countdown_complete', {
                'message': "üß≠ Polaris ready! Click 'Start Bot' to begin."
            })
            print("‚úÖ Countdown complete - Polaris ready!")

    threading.Thread(target=countdown_thread, daemon=True).start()

async def send_greeting():
    if not current_bot_id or bot_startup_state['greeting_sent']:
        return

    try:
        greeting_text = "Hi, I'm Polaris, your meeting AI assistant. Say my name followed by your question to ask me anything!"

        text_hash = hashlib.md5(greeting_text.encode()).hexdigest()
        if text_hash not in audio_cache:
            print("üéµ Generating greeting audio...")

            headers = {
                'api-key': MURF_API_KEY,
                'Content-Type': 'application/json'
            }

            payload = {
                'voiceId': selected_voice_id,
                'text': greeting_text,
                'format': 'mp3',
                'speed': 1.1,
                'pitch': 1.0,
                'style': 'Conversational'
            }

            response = session.post(
                'https://api.murf.ai/v1/speech/generate',
                json=payload,
                headers=headers,
                timeout=12
            )

            if response.status_code == 200:
                content_type = response.headers.get('content-type', '')

                if content_type.startswith('audio/'):
                    audio_data = response.content
                else:
                    result = response.json()
                    if 'audioFile' in result:
                        audio_response = session.get(result['audioFile'], timeout=8)
                        audio_data = audio_response.content
                    elif 'audioContent' in result:
                        audio_data = base64.b64decode(result['audioContent'])
                    else:
                        print("‚ùå Unknown greeting audio format")
                        return

                audio_cache[text_hash] = audio_data

        if text_hash in audio_cache:
            audio_base64 = base64.b64encode(audio_cache[text_hash]).decode('utf-8')

            url = f"https://us-west-2.recall.ai/api/v1/bot/{current_bot_id}/output_audio/"
            headers = {
                'Authorization': f'Token {RECALL_API_KEY}',
                'Content-Type': 'application/json',
                'Accept': 'application/json'
            }

            payload = {
                "kind": "mp3",
                "b64_data": audio_base64
            }

            response = session.post(url, headers=headers, json=payload, timeout=8)

            if response.status_code == 200:
                print("‚úÖ Greeting sent!")
                bot_startup_state['greeting_sent'] = True
                socketio.emit('greeting_sent', {
                })
            else:
                print(f"‚ùå Failed to send greeting: {response.text}")

    except Exception as e:
        print(f"‚ùå Greeting error: {e}")

class RAGOptimizedPipeline:
    def __init__(self):
        self.executor = concurrent.futures.ThreadPoolExecutor(max_workers=3)

    async def process_polaris_request(self, text):
        """COMPLETELY FIXED: RAG-enhanced processing with proper document retrieval"""
        try:
            # Check if muted
            global polaris_muted
            if polaris_muted:
                print(f"üîá Polaris is muted - ignoring: {text}")
                return

            if not check_polaris_mention(text):
                print(f"üö´ 'Polaris' not found in first 3 words: {text}")
                return

            if not bot_startup_state['is_ready']:
                print("üö´ Bot not ready - ignoring transcript")
                return

            start_time = time.time()
            print(f"üß≠ Processing Polaris request: {text}")

            # Add to conversation memory
            global conversation_memory
            conversation_memory.append({
                'timestamp': datetime.now().isoformat(),
                'type': 'user_question',
                'text': text
            })

            # Step 1: RAG search - CRITICAL FIX
            rag_start = time.time()
            relevant_chunks = await self.parallel_rag_search(text)
            rag_time = time.time() - rag_start

            print(f"üîç RAG search ({rag_time:.1f}s): Found {len(relevant_chunks)} chunks")

            # Step 2: AI response with RAG context
            ai_start = time.time()
            ai_response = await self.get_context_aware_response_with_rag(text, relevant_chunks)
            if not ai_response:
                print("‚ùå No AI response")
                return

            conversation_memory.append({
                'timestamp': datetime.now().isoformat(),
                'type': 'ai_response',
                'text': ai_response
            })

            ai_time = time.time() - ai_start
            step1_time = time.time() - start_time

            print(f"üß† AI response ({ai_time:.1f}s): {ai_response}")

            socketio.emit('ai_response', {
                'original_text': text,
                'ai_response': ai_response,
                'processing_time': step1_time,
                'rag_chunks': len(relevant_chunks),
                'mode': response_settings['mode'],
                'rag_search_time': rag_time,
                'memory_items': len(conversation_memory),
                'rag_sources': list(set(chunk['metadata']['filename'] for chunk in relevant_chunks)) if relevant_chunks else []
            })

            # Step 3: TTS
            audio_start = time.time()
            audio_base64 = await self.get_audio_fast(ai_response)
            if not audio_base64:
                print("‚ùå No audio generated")
                return

            step2_time = time.time() - audio_start
            print(f"üéµ Audio generated ({step2_time:.1f}s)")

            # Step 4: Send audio
            send_start = time.time()
            success = await self.send_audio_fast(audio_base64)
            step3_time = time.time() - send_start
            total_time = time.time() - start_time

            if success:
                print(f"‚úÖ TOTAL: {total_time:.1f}s")
                socketio.emit('audio_sent', {
                    'status': 'success',
                    'text': ai_response,
                    'total_time': total_time,
                    'mode': response_settings['mode'],
                    'breakdown': {
                        'rag_time': rag_time,
                        'ai_time': ai_time,
                        'tts_time': step2_time,
                        'send_time': step3_time
                    },
                    'rag_used': len(relevant_chunks) > 0
                })
            else:
                print("‚ùå Failed to send audio")

        except Exception as e:
            print(f"‚ùå Pipeline error: {e}")

    async def parallel_rag_search(self, text):
        def search_task():
            return rag_system.search(text, top_k=5)  # Increased to 5 for better context

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, search_task)

    async def get_context_aware_response_with_rag(self, text, relevant_chunks):
        """COMPLETELY FIXED: Enhanced AI response with proper RAG integration"""
        try:
            def ai_call():
                headers = {
                    'Authorization': f'Bearer {CEREBRAS_API_KEY}',
                    'Content-Type': 'application/json'
                }

                # Build meeting context
                meeting_context = f"You are Polaris, an AI meeting assistant currently participating in this meeting. "

                # Build memory context
                memory_context = ""
                if conversation_memory:
                    memory_context = "Meeting conversation so far:\n"
                    recent_memory = conversation_memory[-6:]
                    for entry in recent_memory:
                        role = "Participant" if entry['type'] == 'user_question' else "You (Polaris)"
                        memory_context += f"{role}: {entry['text'][:60]}...\n"
                    memory_context += "\n"

                # Build document context - CRITICAL FIX
                context = ""
                system_prompt = ""

                # Enhanced RAG context building
                if relevant_chunks:
                    context = "Relevant information from uploaded documents:\n"
                    sources = set()
                    for i, chunk in enumerate(relevant_chunks[:3], 1):  # Top 3 chunks
                        chunk_text = chunk['chunk'][:300]  # More context
                        score = chunk['score']
                        filename = chunk['metadata']['filename']
                        sources.add(filename)
                        context += f"{i}. From {filename} (Score: {score:.3f}): {chunk_text}...\n"
                    context += f"\nSources: {', '.join(sources)}\n\n"

                    if response_settings['mode'] == 'rag' or len(relevant_chunks) > 0:
                        system_prompt = f"{meeting_context}{memory_context}{context}You are in a meeting context. When participants mention you by name, respond helpfully using the document information provided above when relevant. Be conversational and helpful. Answer in 25-45 words and reference the source when using document information."
                        response_settings['max_tokens'] = 90
                    else:
                        system_prompt = f"{meeting_context}{memory_context}You're in a meeting context. Respond helpfully and informatively in 30-45 words."
                        response_settings['max_tokens'] = 85
                else:
                    # No relevant documents found
                    if response_settings['mode'] == 'short':
                        system_prompt = f"{meeting_context}{memory_context}You're in a meeting. Respond helpfully in 15-25 words."
                        response_settings['max_tokens'] = 45
                    elif response_settings['mode'] == 'detailed':
                        system_prompt = f"{meeting_context}{memory_context}You're in a meeting context. Respond helpfully and informatively in 30-45 words."
                        response_settings['max_tokens'] = 85
                    else:
                        system_prompt = f"{meeting_context}{memory_context}You're in a meeting. Use your knowledge to respond helpfully in 25-40 words."
                        response_settings['max_tokens'] = 80

                payload = {
                    'model': 'llama3.1-8b',
                    'messages': [
                        {
                            'role': 'system',
                            'content': system_prompt
                        },
                        {
                            'role': 'user',
                            'content': text
                        }
                    ],
                    'max_tokens': response_settings['max_tokens'],
                    'temperature': 0.4,
                    'top_p': 0.8
                }

                response = session.post(
                    'https://api.cerebras.ai/v1/chat/completions',
                    json=payload,
                    headers=headers,
                    timeout=8
                )

                if response.status_code == 200:
                    result = response.json()
                    full_response = result['choices'][0]['message']['content'].strip()

                    # Don't reference external content in meeting context
                    if "65daysofstatic" in full_response or "song" in full_response.lower() or "album" in full_response.lower():
                        return "I'm here to help with your meeting. What can I assist you with?"

                    # Ensure proper endings
                    if not full_response.endswith(('.', '!', '?', ':')):
                        if len(full_response) > 10:
                            full_response += "."

                    return full_response
                else:
                    print(f"‚ùå AI error: {response.status_code}")
                    return None

            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(self.executor, ai_call)

        except Exception as e:
            print(f"‚ùå AI error: {e}")
            return None

    async def get_audio_fast(self, text):
        try:
            text_hash = hashlib.md5(text.encode()).hexdigest()
            if text_hash in audio_cache:
                print("‚ö° Cache hit!")
                return base64.b64encode(audio_cache[text_hash]).decode('utf-8')

            def audio_call():
                headers = {
                    'api-key': MURF_API_KEY,
                    'Content-Type': 'application/json'
                }

                payload = {
                    'voiceId': selected_voice_id,
                    'text': text,
                    'format': 'mp3',
                    'speed': 1.1,
                    'pitch': 1.0,
                    'style': 'Conversational'
                }

                response = session.post(
                    'https://api.murf.ai/v1/speech/generate',
                    json=payload,
                    headers=headers,
                    timeout=12
                )

                if response.status_code == 200:
                    content_type = response.headers.get('content-type', '')

                    if content_type.startswith('audio/'):
                        audio_data = response.content
                    else:
                        result = response.json()
                        if 'audioFile' in result:
                            audio_response = session.get(result['audioFile'], timeout=8)
                            audio_data = audio_response.content
                        elif 'audioContent' in result:
                            audio_data = base64.b64decode(result['audioContent'])
                        else:
                            return None

                    audio_cache[text_hash] = audio_data
                    return base64.b64encode(audio_data).decode('utf-8')
                else:
                    print(f"‚ùå Audio error: {response.status_code}")
                    return None

            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(self.executor, audio_call)

        except Exception as e:
            print(f"‚ùå Audio error: {e}")
            return None

    async def send_audio_fast(self, base64_audio):
        try:
            if not current_bot_id:
                return False

            def send_call():
                url = f"https://us-west-2.recall.ai/api/v1/bot/{current_bot_id}/output_audio/"

                headers = {
                    'Authorization': f'Token {RECALL_API_KEY}',
                    'Content-Type': 'application/json',
                    'Accept': 'application/json'
                }

                audio_size = len(base64_audio)
                if audio_size > 1000000:
                    base64_audio_truncated = base64_audio[:1000000]
                else:
                    base64_audio_truncated = base64_audio

                payload = {
                    "kind": "mp3",
                    "b64_data": base64_audio_truncated
                }

                response = session.post(url, headers=headers, json=payload, timeout=8)

                if response.status_code == 200:
                    return True
                else:
                    print(f"‚ùå Send failed: {response.text}")
                    return False

            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(self.executor, send_call)

        except Exception as e:
            print(f"‚ùå Send error: {e}")
            return False

# Initialize pipeline
pipeline = RAGOptimizedPipeline()

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# COMPLETE UI with document management
POLARIS_INTERFACE = """
<!DOCTYPE html>
<html>
<head>
    <title>Polaris AI - RAG Enhanced</title>
    <meta charset="UTF-8">
    <style>
        body {
            margin: 0;
            padding: 0;
            font-family: 'Arial', sans-serif;
            background: url('/static/Polaris-logo.jpg') no-repeat center center fixed;
            background-size: cover;
        }
        
        body::before {
            content: "";
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: inherit;
            filter: blur(25px) brightness(0.6);
            z-index: -1;
        }

        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: rgba(60, 0, 100, 0.15);
            padding: 40px;
            border-radius: 20px;
            backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 50, 200, 0.3);
            box-shadow: 0 20px 40px rgba(255, 0, 150, 0.2);
        }

        .header {
            text-align: center;
            margin-bottom: 40px;
            position: relative;
        }

        .polaris-star {
            font-size: 60px;
            color: #ff32c8;
            margin-bottom: 20px;
            text-shadow: 0 0 30px rgba(255, 50, 200, 0.8);
            animation: starGlow 3s ease-in-out infinite alternate;
        }

        @keyframes starGlow {
            0% { text-shadow: 0 0 30px rgba(255, 50, 200, 0.8); }
            100% { text-shadow: 0 0 50px rgba(255, 50, 200, 1), 0 0 80px rgba(255, 0, 150, 0.6); }
        }

        .brand {
            font-size: 48px;
            font-weight: bold;
            color: #ff32c8;
            letter-spacing: 6px;
            margin-bottom: 15px;
            text-shadow: 0 0 30px rgba(255, 50, 200, 0.6);
        }

        .subtitle {
            font-size: 18px;
            color: #ff96e6;
            margin-bottom: 10px;
            font-weight: 300;
        }

        .powered-by {
            font-size: 14px;
            color: #d896e6;
            font-style: italic;
            opacity: 0.8;
        }

        .polaris-note {
            background: linear-gradient(135deg, rgba(255, 50, 200, 0.15), rgba(200, 0, 150, 0.15));
            padding: 20px;
            border-radius: 15px;
            margin: 30px 0;
            border-left: 5px solid #ff32c8;
            text-align: center;
            box-shadow: 0 10px 25px rgba(255, 50, 200, 0.1);
        }

        .polaris-note strong {
            color: #ff32c8;
            font-size: 16px;
        }

        .startup-panel {
            background: linear-gradient(135deg, rgba(255, 0, 150, 0.2), rgba(200, 0, 100, 0.2));
            padding: 30px;
            border-radius: 15px;
            margin: 25px 0;
            border: 2px solid rgba(255, 50, 200, 0.4);
            text-align: center;
            display: none;
            box-shadow: 0 15px 30px rgba(255, 0, 150, 0.15);
        }

        .startup-panel.active {
            display: block;
        }

        .startup-panel h3 {
            color: #ff32c8;
            margin: 0 0 20px 0;
            font-size: 28px;
            text-shadow: 0 0 20px rgba(255, 50, 200, 0.6);
        }

        .countdown-circle {
            width: 100px;
            height: 100px;
            border: 6px solid rgba(255, 50, 200, 0.3);
            border-top: 6px solid #ff32c8;
            border-radius: 50%;
            margin: 20px auto;
            animation: spin 2s linear infinite;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        .countdown-number {
            font-size: 42px;
            font-weight: bold;
            color: #ff32c8;
            margin: 20px 0;
            text-shadow: 0 0 20px rgba(255, 50, 200, 0.8);
        }

        .countdown-message {
            font-size: 16px;
            color: #ffb3e6;
            margin: 15px 0;
        }

        button {
            background: linear-gradient(135deg, #ff32c8, #c800a0);
            color: white;
            border: none;
            padding: 15px 30px;
            border-radius: 12px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            margin: 10px;
            transition: all 0.3s ease;
            box-shadow: 0 8px 20px rgba(255, 50, 200, 0.3);
            text-transform: uppercase;
            letter-spacing: 1px;
        }

        #voice-select {
            width: 100%;
            padding: 12px;
            border: none;
            border-radius: 10px;
            background-color: #2a0040;
            color: #ff96e6;
            font-size: 16px;
            outline: none;
            margin-top: 6px;
            transition: background-color 0.2s, box-shadow 0.2s;
            appearance: none;
            cursor: pointer;
            border: 2px solid rgba(255, 50, 200, 0.3);
        }

        button:hover {
            transform: translateY(-3px);
            box-shadow: 0 12px 30px rgba(255, 50, 200, 0.4);
            background: linear-gradient(135deg, #c800a0, #ff32c8);
        }

        button:disabled {
            opacity: 0.5;
            cursor: not-allowed;
            transform: none;
            background: #662d66;
        }

        .start-bot-button {
            background: linear-gradient(135deg, #9932cc, #6a1b9a);
            box-shadow: 0 8px 20px rgba(153, 50, 204, 0.3);
            display: none;
        }

        .start-bot-button.ready {
            display: inline-block;
        }

        .start-bot-button:hover {
            background: linear-gradient(135deg, #6a1b9a, #9932cc);
            box-shadow: 0 12px 30px rgba(153, 50, 204, 0.4);
        }

        .leave-bot-button {
            background: linear-gradient(135deg, #e91e63, #ad1457);
            box-shadow: 0 8px 20px rgba(233, 30, 99, 0.3);
            display: none;
        }

        .leave-bot-button.ready {
            display: inline-block;
        }

        .leave-bot-button:hover {
            background: linear-gradient(135deg, #ad1457, #e91e63);
            box-shadow: 0 12px 30px rgba(233, 30, 99, 0.4);
        }

        .download-mom-button {
            background: linear-gradient(135deg, #8e24aa, #4a148c);
            box-shadow: 0 8px 20px rgba(142, 36, 170, 0.3);
            display: none;
        }

        .download-mom-button.ready {
            display: inline-block;
        }

        .download-mom-button:hover {
            background: linear-gradient(135deg, #4a148c, #8e24aa);
            box-shadow: 0 12px 30px rgba(142, 36, 170, 0.4);
        }

        .mute-button {
            background: linear-gradient(135deg, #ff1493, #c71585);
            box-shadow: 0 8px 20px rgba(255, 20, 147, 0.3);
            font-size: 18px;
            padding: 12px 25px;
        }

        .mute-button.muted {
            background: linear-gradient(135deg, #666, #888);
            box-shadow: 0 8px 20px rgba(102, 102, 102, 0.3);
        }

        .mute-button:hover {
            background: linear-gradient(135deg, #c71585, #ff1493);
            box-shadow: 0 12px 30px rgba(255, 20, 147, 0.4);
        }

        .mute-button.muted:hover {
            background: linear-gradient(135deg, #888, #666);
            box-shadow: 0 12px 30px rgba(102, 102, 102, 0.4);
        }

        .delete-btn {
            background: linear-gradient(135deg, #e91e63, #ad1457);
            padding: 5px 10px;
            font-size: 12px;
            margin: 0 5px;
            text-transform: none;
            letter-spacing: 0;
        }

        .delete-btn:hover {
            background: linear-gradient(135deg, #ad1457, #e91e63);
        }

        input[type="url"], input[type="text"] {
            width: 100%;
            padding: 15px 20px;
            border: 2px solid rgba(255, 50, 200, 0.3);
            border-radius: 12px;
            background: rgba(60, 0, 100, 0.3);
            color: #ff96e6;
            font-size: 16px;
            box-sizing: border-box;
            margin: 10px 0;
            transition: all 0.3s ease;
        }

        input[type="url"]:focus, input[type="text"]:focus {
            outline: none;
            border-color: #ff32c8;
            box-shadow: 0 0 20px rgba(255, 50, 200, 0.4);
            background: rgba(60, 0, 100, 0.5);
        }

        input[type="url"]::placeholder, input[type="text"]::placeholder {
            color: rgba(255, 150, 230, 0.6);
        }

        .mode-selector {
            display: flex;
            justify-content: center;
            margin: 30px 0;
            gap: 15px;
            flex-wrap: wrap;
        }

        .mode-btn {
            background: rgba(255, 50, 200, 0.2);
            color: #ff96e6;
            border: 2px solid rgba(255, 50, 200, 0.3);
            padding: 10px 20px;
            border-radius: 25px;
            font-size: 14px;
            cursor: pointer;
            transition: all 0.3s ease;
            text-transform: uppercase;
            letter-spacing: 1px;
            font-weight: 500;
        }

        .mode-btn:hover, .mode-btn.active {
            background: rgba(255, 50, 200, 0.3);
            border-color: #ff32c8;
            color: #ff32c8;
            box-shadow: 0 0 15px rgba(255, 50, 200, 0.3);
        }

        .upload-area {
            border: 3px dashed rgba(255, 50, 200, 0.4);
            border-radius: 15px;
            padding: 40px;
            text-align: center;
            margin: 30px 0;
            cursor: pointer;
            transition: all 0.3s ease;
            background: rgba(255, 50, 200, 0.05);
        }

        .upload-area:hover {
            border-color: #ff32c8;
            background: rgba(255, 50, 200, 0.1);
            box-shadow: 0 0 30px rgba(255, 50, 200, 0.2);
        }

        .upload-area.dragover {
            border-color: #9932cc;
            background: rgba(153, 50, 204, 0.1);
        }

        .file-input {
            display: none;
        }

        .document-list {
            margin-top: 20px;
            max-height: 200px;
            overflow-y: auto;
            background: rgba(60, 0, 100, 0.3);
            border-radius: 10px;
            padding: 15px;
        }

        .document-item {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 10px;
            margin: 5px 0;
            background: rgba(255, 50, 200, 0.1);
            border-radius: 8px;
            border-left: 3px solid #ff32c8;
        }

        .document-info {
            flex-grow: 1;
        }

        .document-name {
            font-weight: bold;
            color: #ff32c8;
            font-size: 14px;
        }

        .document-details {
            font-size: 11px;
            color: #ff96e6;
            margin-top: 2px;
        }

        .grid-2 {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 30px;
            margin: 30px 0;
        }

        .card {
            background: rgba(60, 0, 100, 0.3);
            padding: 25px;
            border-radius: 15px;
            border: 1px solid rgba(255, 50, 200, 0.2);
            backdrop-filter: blur(10px);
            transition: all 0.3s ease;
        }

        .card:hover {
            border-color: rgba(255, 50, 200, 0.4);
            box-shadow: 0 10px 30px rgba(255, 50, 200, 0.1);
        }

        .card h3 {
            color: #ff32c8;
            margin-top: 0;
            font-size: 20px;
            text-shadow: 0 0 10px rgba(255, 50, 200, 0.5);
        }

        .status-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin: 25px 0;
        }

        .status-item {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 15px 20px;
            background: rgba(60, 0, 100, 0.4);
            border-radius: 12px;
            border-left: 4px solid #ff32c8;
            color: #ff96e6;
        }

        .status-indicator {
            width: 14px;
            height: 14px;
            border-radius: 50%;
            margin-left: 10px;
            animation: pulse 2s infinite;
        }

        @keyframes pulse {
            0% { opacity: 1; }
            50% { opacity: 0.5; }
            100% { opacity: 1; }
        }

        .status-ready { background: #9932cc; }
        .status-active { background: #ff32c8; }
        .status-connecting { background: #ff1493; }
        .status-error { background: #e91e63; }
        .status-muted { background: #666; }

        .log-panel {
            background: rgba(20, 0, 40, 0.8);
            padding: 25px;
            border-radius: 15px;
            margin: 30px 0;
            max-height: 300px;
            overflow-y: auto;
            font-family: 'Courier New', monospace;
            font-size: 13px;
            border: 1px solid rgba(255, 50, 200, 0.2);
        }

        .log-entry {
            margin: 8px 0;
            padding: 5px 0;
            border-bottom: 1px solid rgba(255, 50, 200, 0.1);
            color: #ff96e6;
        }

        .log-entry:last-child {
            border-bottom: none;
        }

        .log-transcript {
            color: #ffb3e6;
            font-style: italic;
        }

        .log-polaris {
            color: #9932cc;
            font-weight: bold;
        }

        .log-ai {
            color: #ff1493;
        }

        .log-rag {
            color: #00ff88;
            font-weight: bold;
        }

        .log-memory {
            color: #e91e63;
        }

        .log-mom {
            color: #8e24aa;
            font-weight: bold;
        }

        .log-mute {
            color: #ff1493;
            font-weight: bold;
        }

        @media (max-width: 768px) {
            .container { padding: 20px; }
            .grid-2 { grid-template-columns: 1fr; }
            .brand { font-size: 36px; }
            .polaris-star { font-size: 40px; }
        }
    </style>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/socket.io/4.0.1/socket.io.js"></script>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="brand">POLARIS</div>
            <div class="subtitle">RAG-Enhanced AI Meeting Assistant</div>
            <div class="powered-by">üéØ Powered By: Murf AI, Cerebras, Recall.ai + RAG System</div>
        </div>

        <div class="polaris-note">
            <div class="input-group">
                <label class="subtitle">Assistant Voice: </label>
                <select id="voice-select" onchange="changeVoice()">
                    <!-- Options populated by JavaScript -->
                </select>
            </div>
        </div>

        <!-- Mute/Unmute Button -->
        <div style="text-align: center; margin: 20px 0;">
            <button class="mute-button" id="mute-button" onclick="toggleMute()">
                üîä POLARIS ACTIVE
            </button>
        </div>

        <!-- Startup Panel -->
        <div class="startup-panel" id="startup-panel">
            <h3>üß≠ Polaris Initialization</h3>
            <div class="countdown-circle" id="countdown-circle"></div>
            <div class="countdown-number" id="countdown-number">--</div>
            <div class="countdown-message" id="countdown-message">Polaris joining meeting...</div>
            <button class="start-bot-button" id="start-bot-button" onclick="startBot()">
                üöÄ Start Polaris
            </button>
            <button class="leave-bot-button" id="leave-bot-button" onclick="leaveBot()">
                üö™ Leave Meeting
            </button>
        </div>

        <!-- Response Mode Selector -->
        <div class="mode-selector">
            <button class="mode-btn active" data-mode="detailed" onclick="setMode('detailed')">üîç Detailed</button>
            <button class="mode-btn" data-mode="short" onclick="setMode('short')">‚ö° Quick</button>
            <button class="mode-btn" data-mode="rag" onclick="setMode('rag')">üìö Document</button>
            <button class="mode-btn" data-mode="internet" onclick="setMode('internet')">üåê General</button>
        </div>

        <div class="grid-2">
            <!-- Meeting Controls -->
            <div class="card">
                <h3>üîó Meeting Connection</h3>
                <input type="url" id="meeting-url" placeholder="Enter meeting URL...">
                <div style="text-align: center;">
                    <button onclick="startTunnel()" id="tunnel-btn">üöá Start Tunnel</button>
                    <button onclick="joinMeeting()" id="join-btn" disabled>üß≠ Join Meeting</button>
                </div>
                <div style="text-align: center; margin-top: 15px;">
                    <button onclick="leaveBot()" id="leave-meeting-btn" class="leave-bot-button">
                        üö™ Leave Meeting
                    </button>
                    <button onclick="downloadMOM()" id="download-mom-btn" class="download-mom-button">
                        üìÑ Download MOM
                    </button>
                </div>
            </div>

            <!-- Enhanced Document Upload with Management -->
            <div class="card">
                <h3>üìö RAG Document Management</h3>
                <div class="upload-area" onclick="document.getElementById('file-input').click()">
                    <div style="font-size: 24px; margin-bottom: 10px;">üìÅ</div>
                    <div style="color: white;"><strong>Upload documents for RAG context</strong></div>
                    <div style="font-size: 12px; color: #d896e6; margin-top: 8px;">
                        Supports: PDF, DOCX, TXT (Max 16MB)<br>
                        Documents will be used to answer questions contextually
                    </div>
                    <input type="file" id="file-input" class="file-input" multiple accept=".pdf,.docx,.doc,.txt">
                </div>

                <div class="document-list" id="document-list" style="display: none;">
                    <h4 style="color: #ff32c8; margin: 0 0 15px 0;">üìã Uploaded Documents:</h4>
                    <div id="document-items"></div>
                </div>

                <div style="text-align: center; margin-top: 15px;">
                    <span style="font-size: 14px; color: #ff96e6;">
                        Documents: <span id="doc-count">0</span> | 
                        Chunks: <span id="chunk-count">0</span> |
                        RAG Status: <span id="rag-status" style="color: #00ff88;">Ready</span>
                    </span>
                </div>
            </div>
        </div>

        <!-- System Status -->
        <div class="card">
            <h3>üìä System Status</h3>
            <div class="status-grid">
                <div class="status-item">
                    <span>Tunnel Connection</span>
                    <span>
                        <span id="tunnel-status">Not Started</span>
                        <span class="status-indicator status-ready" id="tunnel-indicator"></span>
                    </span>
                </div>
                <div class="status-item">
                    <span>Polaris Bot</span>
                    <span>
                        <span id="bot-status">Inactive</span>
                        <span class="status-indicator status-ready" id="bot-indicator"></span>
                    </span>
                </div>
                <div class="status-item">
                    <span>RAG System</span>
                    <span>
                        <span id="rag-system-status">Ready</span>
                        <span class="status-indicator status-ready" id="rag-system-indicator"></span>
                    </span>
                </div>
                <div class="status-item">
                    <span>Response Mode</span>
                    <span>
                        <span id="current-mode">Detailed</span>
                        <span class="status-indicator status-active" id="mode-indicator"></span>
                    </span>
                </div>
                <div class="status-item">
                    <span>Polaris Status</span>
                    <span>
                        <span id="mute-status">Active</span>
                        <span class="status-indicator status-active" id="mute-indicator"></span>
                    </span>
                </div>
            </div>
        </div>

        <!-- Enhanced Activity Log -->
        <div class="card">
            <h3>üìã Live Meeting Transcript & RAG Activity Log</h3>
            <div class="log-panel" id="log-panel">
                <div class="log-entry">üß≠ Polaris RAG-Enhanced AI Meeting Assistant Started</div>
                <div class="log-entry log-rag">üìö RAG system initialized and ready for document uploads</div>
            </div>
        </div>
    </div>

    <script>
        const socket = io();
        let uploadedDocs = [];
        let currentMode = 'detailed';
        let momFilename = null;
        let polarisIsMuted = false;

        function updateStatus(type, status, indicator) {
            document.getElementById(type + '-status').textContent = status;
            document.getElementById(type + '-indicator').className = 'status-indicator ' + indicator;
        }

        function addLog(message, type = 'info') {
            const logPanel = document.getElementById('log-panel');
            const logEntry = document.createElement('div');
            logEntry.className = 'log-entry';

            if (type === 'transcript') {
                logEntry.classList.add('log-transcript');
            } else if (type === 'polaris') {
                logEntry.classList.add('log-polaris');
            } else if (type === 'ai') {
                logEntry.classList.add('log-ai');
            } else if (type === 'rag') {
                logEntry.classList.add('log-rag');
            } else if (type === 'memory') {
                logEntry.classList.add('log-memory');
            } else if (type === 'mom') {
                logEntry.classList.add('log-mom');
            } else if (type === 'mute') {
                logEntry.classList.add('log-mute');
            }

            logEntry.textContent = new Date().toLocaleTimeString() + ' - ' + message;
            logPanel.appendChild(logEntry);
            logPanel.scrollTop = logPanel.scrollHeight;
        }

        // Mute/Unmute functionality
        function toggleMute() {
            polarisIsMuted = !polarisIsMuted;
            const muteButton = document.getElementById('mute-button');
            const muteStatus = document.getElementById('mute-status');
            const muteIndicator = document.getElementById('mute-indicator');

            if (polarisIsMuted) {
                muteButton.textContent = 'üîá POLARIS MUTED';
                muteButton.classList.add('muted');
                muteStatus.textContent = 'Muted';
                muteIndicator.className = 'status-indicator status-muted';
                addLog('üîá Polaris responses muted.', 'mute');
            } else {
                muteButton.textContent = 'üîä POLARIS ACTIVE';
                muteButton.classList.remove('muted');
                muteStatus.textContent = 'Active';
                muteIndicator.className = 'status-indicator status-active';
                addLog('üîä Polaris responses enabled - will respond to "Polaris" mentions', 'mute');
            }

            socket.emit('set_mute_state', { muted: polarisIsMuted });
        }

        function setMode(mode) {
            currentMode = mode;
            document.querySelectorAll('.mode-btn').forEach(btn => {
                btn.classList.remove('active');
            });
            document.querySelector(`[data-mode="${mode}"]`).classList.add('active');
            document.getElementById('current-mode').textContent = mode.charAt(0).toUpperCase() + mode.slice(1);

            socket.emit('set_response_mode', { mode: mode });
            addLog(`Response mode changed to: ${mode}`);
        }

        function showStartupSequence() {
            document.getElementById('startup-panel').classList.add('active');
            document.getElementById('leave-bot-button').classList.add('ready');
            document.getElementById('leave-meeting-btn').classList.add('ready');
            updateStatus('bot', 'Joining...', 'status-connecting');
        }

        function hideStartupSequence() {
            document.getElementById('startup-panel').classList.remove('active');
        }

        function startBot() {
            updateStatus('bot', 'Active', 'status-active');
            hideStartupSequence();
            socket.emit('start_bot_manual');
        }

        function leaveBot() {
            if (confirm('Are you sure you want Polaris to leave the meeting? This will generate professional MOM.')) {
                socket.emit('leave_bot');
            }
        }

        function downloadMOM() {
            if (!momFilename) {
                alert('No Meeting Minutes available for download yet.');
                return;
            }

            const downloadUrl = `/download_mom/${momFilename}`;
            const link = document.createElement('a');
            link.href = downloadUrl;
            link.download = momFilename;
            document.body.appendChild(link);
            link.click();
            document.body.removeChild(link);
        }

        function updateDocumentList(documents) {
            const documentList = document.getElementById('document-list');
            const documentItems = document.getElementById('document-items');

            if (documents.length > 0) {
                documentList.style.display = 'block';
                documentItems.innerHTML = '';

                documents.forEach(doc => {
                    const docItem = document.createElement('div');
                    docItem.className = 'document-item';
                    docItem.innerHTML = `
                        <div class="document-info">
                            <div class="document-name">üìÑ ${doc.filename}</div>
                            <div class="document-details">${doc.chunk_count} chunks ‚Ä¢ ${doc.total_chars} chars ‚Ä¢ Uploaded: ${new Date(doc.upload_time).toLocaleTimeString()}</div>
                        </div>
                        <button class="delete-btn" onclick="removeDocument('${doc.filename}')">
                            ‚úñ Delete
                        </button>
                    `;
                    documentItems.appendChild(docItem);
                });
            } else {
                documentList.style.display = 'none';
            }
        }

        function removeDocument(filename) {
            if (confirm(`Are you sure you want to remove "${filename}" from RAG context?`)) {
                addLog(`üóëÔ∏è Removing document from RAG: ${filename}`, 'rag');
                socket.emit('remove_document', { filename: filename });
            }
        }

        // Enhanced file upload handling
        document.getElementById('file-input').addEventListener('change', function(e) {
            uploadFiles(e.target.files);
        });

        const uploadArea = document.querySelector('.upload-area');
        uploadArea.addEventListener('dragover', function(e) {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });

        uploadArea.addEventListener('dragleave', function(e) {
            uploadArea.classList.remove('dragover');
        });

        uploadArea.addEventListener('drop', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            uploadFiles(e.dataTransfer.files);
        });

        function uploadFiles(files) {
            Array.from(files).forEach(file => {
                const formData = new FormData();
                formData.append('file', file);
                addLog('üì§ Processing for RAG: ' + file.name, 'rag');

                fetch('/upload', {
                    method: 'POST',
                    body: formData
                })
                .then(response => response.json())
                .then(data => {
                    if (data.success) {
                        addLog(`‚úÖ Added to RAG: ${file.name} (${data.chunks} chunks)`, 'rag');
                        updateDocumentStats();
                        socket.emit('get_document_list');
                        updateStatus('rag-system', 'Active', 'status-active');
                    } else {
                        addLog(`‚ùå Failed to process for RAG: ${file.name}`, 'rag');
                    }
                })
                .catch(error => {
                    addLog(`‚ùå Upload error: ${error}`, 'rag');
                });
            });
        }

        function updateDocumentStats() {
            socket.emit('get_rag_stats');
        }

        function startTunnel() {
            updateStatus('tunnel', 'Starting...', 'status-connecting');
            addLog('Starting secure tunnel...');
            socket.emit('start_tunnel');
        }

        function joinMeeting() {
            const meetingUrl = document.getElementById('meeting-url').value.trim();
            if (!meetingUrl) {
                alert('Please enter a meeting URL');
                return;
            }
            updateStatus('bot', 'Creating...', 'status-connecting');
            addLog('Creating RAG-enhanced Polaris bot...');
            socket.emit('join_meeting', { meetingUrl: meetingUrl });
        }

        // Socket event handlers
        socket.on('tunnel_ready', function(data) {
            updateStatus('tunnel', 'Connected', 'status-active');
            document.getElementById('join-btn').disabled = false;
            addLog('Tunnel connection established');
        });

        socket.on('bot_created', function(data) {
            updateStatus('bot', 'Created', 'status-active');
            addLog('üß≠ RAG-enhanced Polaris bot created: ' + data.botId);
            showStartupSequence();
        });

        socket.on('bot_error', function(data) {
            updateStatus('bot', 'Error', 'status-error');
            addLog('‚ùå Bot error: ' + data.error);
        });

        socket.on('bot_left', function(data) {
            updateStatus('bot', 'Left Meeting', 'status-ready');
            document.getElementById('leave-bot-button').classList.remove('ready');
            document.getElementById('leave-meeting-btn').classList.remove('ready');
            addLog('üö™ ' + data.message, 'mom');
            hideStartupSequence();
        });

        socket.on('bot_leave_error', function(data) {
            updateStatus('bot', 'Leave Failed', 'status-error');
            addLog('‚ùå Leave failed: ' + data.error);
        });

        socket.on('mom_generated', function(data) {
            addLog('üìù ' + data.message, 'mom');
            addLog('üìÑ MOM generated successfully!', 'mom');

            momFilename = data.pdf_filename;
            document.getElementById('download-mom-btn').classList.add('ready');
        });

        socket.on('transcript_received', function(data) {
            addLog('üìù Transcript: ' + data.text, 'transcript');
        });

        socket.on('countdown_update', function(data) {
            document.getElementById('countdown-number').textContent = data.seconds_remaining;
            document.getElementById('countdown-message').textContent = data.message;
        });

        socket.on('enable_start_button', function(data) {
            document.getElementById('start-bot-button').classList.add('ready');
            addLog('üöÄ Polaris ready to start');
        });

        socket.on('countdown_complete', function(data) {
            document.getElementById('countdown-circle').style.display = 'none';
            document.getElementById('countdown-number').textContent = '‚úÖ';
            document.getElementById('countdown-message').textContent = data.message;
            document.getElementById('start-bot-button').classList.add('ready');
            updateStatus('bot', 'Ready', 'status-ready');
        });

        socket.on('greeting_sent', function(data) {
            addLog('üéâ Polaris greeting sent - RAG system active');
            updateStatus('rag-system', 'Active', 'status-active');
        });

        socket.on('ai_response', function(data) {
            let ragInfo = '';
            if (data.rag_chunks > 0) {
                ragInfo = ` [RAG: ${data.rag_chunks} chunks`;
                if (data.rag_sources && data.rag_sources.length > 0) {
                    ragInfo += ` from ${data.rag_sources.join(', ')}`;
                }
                ragInfo += ']';
            }

            addLog(`üß† AI (${data.mode}): ${data.ai_response}${ragInfo}`, data.rag_chunks > 0 ? 'rag' : 'ai');

            if (data.memory_items) {
                addLog(`üß† Memory updated: ${data.memory_items} items`, 'memory');
            }
        });

        socket.on('audio_sent', function(data) {
            updateStatus('rag-system', 'Ready', 'status-ready');

            const ragUsed = data.rag_used ? ' (RAG Enhanced)' : '';
        });

        socket.on('document_list', function(data) {
            updateDocumentList(data.documents);
        });

        socket.on('document_removed', function(data) {
            addLog(`‚úÖ Removed from RAG: ${data.filename}`, 'rag');
            updateDocumentStats();
            socket.emit('get_document_list');
        });

        socket.on('document_remove_error', function(data) {
            addLog(`‚ùå Remove failed: ${data.error}`, 'rag');
        });

        socket.on('rag_stats', function(data) {
            document.getElementById('doc-count').textContent = data.stats.total_documents;
            document.getElementById('chunk-count').textContent = data.stats.total_chunks;

            if (data.stats.total_chunks > 0) {
                document.getElementById('rag-status').textContent = 'Active';
                document.getElementById('rag-status').style.color = '#00ff88';
            } else {
                document.getElementById('rag-status').textContent = 'Ready';
                document.getElementById('rag-status').style.color = '#ff96e6';
            }
        });

        // Initialize
        document.addEventListener('DOMContentLoaded', function() {
            updateStatus('rag-system', 'Ready', 'status-ready');
            socket.emit('get_document_list');
            socket.emit('get_rag_stats');
        });

        // API key management
        document.addEventListener("DOMContentLoaded", function() {
            fetch("/get_api_keys")
                .then(res => res.json())
                .then(keys => {
                    let missing = [];
                    if (!keys.RECALL_API_KEY) missing.push("RECALL_API_KEY");
                    if (!keys.CEREBRAS_API_KEY) missing.push("CEREBRAS_API_KEY");
                    if (!keys.MURF_API_KEY) missing.push("MURF_API_KEY");
                    if (missing.length > 0) {
                        let newKeys = {};
                        missing.forEach(k => {
                            let val = prompt(`Please enter your ${k}:`);
                            if (val) newKeys[k] = val;
                        });
                        if (Object.keys(newKeys).length > 0) {
                            fetch("/set_api_keys", {
                                method: "POST",
                                headers: { "Content-Type": "application/json" },
                                body: JSON.stringify(newKeys)
                            }).then(() => {
                                alert("‚úÖ API keys saved. Please refresh the page.");
                            });
                        }
                    }
                });
        });

        // Voice options population
        document.addEventListener("DOMContentLoaded", function() {
            const voiceSelect = document.getElementById('voice-select');
            const voices = [
                ["Marcus (US Male)", "en-US-marcus"],
                ["Natalie (US Female)", "en-US-natalie"],
                ["Amara (US Female)", "en-US-amara"],
                ["Charles (US Male)", "en-US-charles"],
                ["Freddie (UK Male)", "en-UK-freddie"],
                ["Emma (UK Female)", "en-UK-emma"],
                ["Oliver (UK Male)", "en-UK-oliver"],
                ["Sarah (US Female)", "en-US-sarah"],
                ["David (UK Male)", "en-UK-david"],
                ["Sophie (UK Female)", "en-UK-sophie"],
                ["Alex (US Male)", "en-US-alex"],
                ["Lily (US Female)", "en-US-lily"],
                ["James (UK Male)", "en-UK-james"],
                ["Grace (UK Female)", "en-UK-grace"],
                ["Ryan (US Male)", "en-US-ryan"]
            ];
            voices.forEach(function(v){
                const opt = document.createElement('option');
                opt.value = v[1];
                opt.textContent = v[0];
                voiceSelect.appendChild(opt);
            });
        });

        function changeVoice() {
            const selected = document.getElementById('voice-select').value;
            socket.emit('set_voice', { voiceId: selected });
            addLog(`üé§ Voice changed to: ${selected}`);
        }
    </script>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(POLARIS_INTERFACE)

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'})

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})

        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': 'File type not supported'})

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        chunks = []
        file_ext = filename.lower().split('.')[-1]

        start_time = time.time()
        if file_ext == 'pdf':
            chunks = rag_system.process_pdf(filepath)
        elif file_ext in ['docx', 'doc']:
            chunks = rag_system.process_docx(filepath)
        elif file_ext == 'txt':
            chunks = rag_system.process_txt(filepath)

        processing_time = time.time() - start_time
        print(f"üìÅ RAG processing time: {processing_time:.2f}s for {len(chunks)} chunks")

        if chunks:
            success = rag_system.add_document(filename, chunks)
            if success:
                stats = rag_system.get_stats()
                return jsonify({
                    'success': True,
                    'filename': filename,
                    'chunks': len(chunks),
                    'processing_time': processing_time,
                    'stats': stats
                })

        return jsonify({'success': False, 'error': 'Failed to process document for RAG'})

    except Exception as e:
        print(f"‚ùå Upload error: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/download_mom/<filename>')
def download_mom(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        print(f"‚ùå Download error: {e}")
        return jsonify({'error': str(e)}), 500

# Enhanced webhook with better transcript handling
@app.route('/webhook', methods=['POST'])
def webhook():
    global transcript_buffer, last_update_time, sentence_timeout_handle, full_meeting_transcript

    try:
        data = request.json
        event = data.get('event')
        transcript_data = data.get('data')

        if event == 'transcript.data':
            transcript_text = ''

            if transcript_data and transcript_data.get('data') and transcript_data['data'].get('words'):
                transcript_text = ' '.join([word['text'] for word in transcript_data['data']['words']]).strip()
            elif transcript_data and transcript_data.get('words'):
                transcript_text = ' '.join([word['text'] for word in transcript_data['words']]).strip()
            elif transcript_data and transcript_data.get('text'):
                transcript_text = transcript_data['text'].strip()

            if transcript_text:
                current_time = time.time() * 1000

                if sentence_timeout_handle:
                    sentence_timeout_handle.cancel()

                if current_time - last_update_time > 4000:
                    if transcript_buffer.strip():
                        finalize_and_emit_transcript()
                    transcript_buffer = transcript_text
                else:
                    if transcript_buffer and not transcript_buffer.endswith(' '):
                        transcript_buffer += ' '
                    transcript_buffer += transcript_text

                last_update_time = current_time

                sentence_timeout_handle = threading.Timer(3.5, finalize_and_emit_transcript)
                sentence_timeout_handle.start()

                if (transcript_buffer and 
                    len(transcript_buffer.split()) >= 6 and
                    (any(transcript_buffer.strip().endswith(punct) for punct in ['. ', '! ', '? ']) or
                     transcript_buffer.strip().endswith('.') or
                     transcript_buffer.strip().endswith('!') or  
                     transcript_buffer.strip().endswith('?'))):
                    if sentence_timeout_handle:
                        sentence_timeout_handle.cancel()
                    finalize_and_emit_transcript()

        return jsonify({'success': True})

    except Exception as e:
        print(f"‚ùå Webhook error: {e}")
        return jsonify({'error': 'Webhook processing failed'}), 500

@app.route('/set_voice', methods=['POST'])
def set_voice():
    global selected_voice_id
    data = request.json
    vid = data.get('voice_id', 'en-US-marcus')

    voice_dict = dict(voice_options)
    if vid not in voice_dict.values():
        return jsonify({'success': False, 'error': 'Invalid voice'})

    selected_voice_id = vid

    for name, value in voice_options:
        if value == vid:
            return jsonify({'success': True, 'voice_id': vid, 'voice_name': name})
    return jsonify({'success': True, 'voice_id': vid, 'voice_name': vid})

def finalize_and_emit_transcript():
    global transcript_buffer

    if transcript_buffer.strip():
        complete_sentence = transcript_buffer.strip()

        print(f"üìù Complete sentence: {complete_sentence}")

        socketio.emit('transcript_received', {'text': complete_sentence})

        full_meeting_transcript.append({
            'timestamp': datetime.now().isoformat(),
            'text': complete_sentence
        })

        def run_pipeline():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            loop.run_until_complete(pipeline.process_polaris_request(complete_sentence))
            loop.close()

        threading.Thread(target=run_pipeline, daemon=True).start()
        transcript_buffer = ''

# Socket handlers
@socketio.on('start_tunnel')
def handle_start_tunnel():
    global tunnel_url

    def start_tunnel_thread():
        global tunnel_url
        try:
            if tunnel_url:
                socketio.emit('tunnel_ready', {'url': tunnel_url})
                return

            print("üöá Starting tunnel...")
            ngrok.kill()
            tunnel_obj = ngrok.connect(5013, "http")

            if hasattr(tunnel_obj, 'public_url'):
                tunnel_url = tunnel_obj.public_url
            else:
                tunnel_url = str(tunnel_obj).split('"')[1] if '"' in str(tunnel_obj) else str(tunnel_obj)

            print(f"‚úÖ Tunnel ready: {tunnel_url}")
            socketio.emit('tunnel_ready', {'url': tunnel_url})

        except Exception as e:
            print(f"‚ùå Tunnel error: {e}")
            socketio.emit('tunnel_error', {'message': str(e)})

    threading.Thread(target=start_tunnel_thread, daemon=True).start()

@socketio.on('join_meeting')
def handle_join_meeting(data):
    global current_bot_id, tunnel_url

    meeting_url = data.get('meetingUrl')
    if not meeting_url or not tunnel_url:
        emit('bot_error', {'error': 'Missing meeting URL or tunnel'})
        return

    def create_polaris_bot():
        global current_bot_id

        webhook_url = f'{tunnel_url}/webhook'
        bot_id, error = create_recall_bot_correct(meeting_url, webhook_url)

        if bot_id:
            current_bot_id = bot_id
            print(f"‚úÖ RAG-enhanced Polaris bot ready: {current_bot_id}")
            socketio.emit('bot_created', {
                'success': True,
                'botId': current_bot_id,
                'message': 'RAG-enhanced Polaris bot created!'
            })
            start_bot_countdown()
        else:
            print(f"‚ùå Bot creation failed: {error}")
            socketio.emit('bot_error', {'error': f'Bot creation failed: {error}'})

    threading.Thread(target=create_polaris_bot, daemon=True).start()

@socketio.on('start_bot_manual')
def handle_start_bot_manual():
    global bot_startup_state

    if not bot_startup_state['is_joining']:
        print("‚ö†Ô∏è Bot start called but not in joining state")
        return

    bot_startup_state['is_joining'] = False
    bot_startup_state['is_ready'] = True

    print("üöÄ RAG-enhanced Polaris manually started!")

    def send_greeting_async():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        loop.run_until_complete(send_greeting())
        loop.close()

    threading.Thread(target=send_greeting_async, daemon=True).start()

@socketio.on('leave_bot')
def handle_leave_bot():
    def leave_bot_thread():
        success, message = leave_bot()
        if success:
            socketio.emit('bot_left', {'message': message})
        else:
            socketio.emit('bot_leave_error', {'error': message})

    threading.Thread(target=leave_bot_thread, daemon=True).start()

@socketio.on('set_response_mode')
def handle_set_response_mode(data):
    mode = data.get('mode', 'detailed')
    response_settings['mode'] = mode
    print(f"üîÑ Response mode changed to: {mode}")

@socketio.on('set_mute_state')
def handle_set_mute_state(data):
    global polaris_muted
    muted = data.get('muted', False)
    polaris_muted = muted
    status = "muted" if muted else "active"
    print(f"üîá Polaris {status}")

# RAG document management socket handlers
@socketio.on('get_document_list')
def handle_get_document_list():
    documents = rag_system.get_document_list()
    emit('document_list', {'documents': documents})

@socketio.on('remove_document')
def handle_remove_document(data):
    filename = data.get('filename')
    if filename:
        success, message = rag_system.remove_document(filename)
        if success:
            emit('document_removed', {'filename': filename, 'message': message})
        else:
            emit('document_remove_error', {'filename': filename, 'error': message})

@socketio.on('get_rag_stats')
def handle_get_rag_stats():
    stats = rag_system.get_stats()
    emit('rag_stats', {'stats': stats})

@socketio.on("set_voice")
def handle_set_voice(data):
    global selected_voice_id
    new_voice = data.get("voiceId")
    if new_voice:
        selected_voice_id = new_voice
        print(f"üé§ Voice changed to: {selected_voice_id}")
        emit("voice_changed", {"voiceId": selected_voice_id}, broadcast=True)

@app.route('/health')
def health():
    return jsonify({
        'status': 'healthy',
        'mode': 'Polaris_RAG_ENHANCED_FIXED',
        'target_latency': '4_6_seconds',
        'response_words': '25_45',
        'bot': 'active' if current_bot_id else 'inactive',
        'bot_ready': bot_startup_state['is_ready'],
        'muted': polaris_muted,
        'memory_items': len(conversation_memory),
        'transcript_entries': len(full_meeting_transcript),
        'mom_file': generated_mom_file,
        'rag_stats': rag_system.get_stats(),
        'features': {
            'rag_enhanced': True,
            'document_processing': True,
            'semantic_search': True,
            'context_awareness': True,
            'meeting_context_prompts': True,
            'mute_unmute_control': True,
            'better_memory': True,
            'random_countdown': '16_19_seconds',
            'faiss_l2_index': True,
            'cosine_similarity': True,
            'lowered_threshold': True
        }
    })

if __name__ == '__main__':
    print("üß≠ POLARIS RAG-ENHANCED AI - STARTED AT http://localhost:5013/")
    print("=" * 80)
    socketio.run(app, host='0.0.0.0', port=5013, debug=False)
